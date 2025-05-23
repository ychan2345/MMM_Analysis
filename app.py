import os
import base64
from openai import OpenAI
import streamlit as st
from PIL import Image, ImageEnhance, ImageOps, ImageFilter
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import numpy as np
import pytesseract
import tempfile
import matplotlib.pyplot as plt
import json
from matplotlib.figure import Figure

def enhance_image(image_bytes: bytes,
                  sharpness_factor=2.0,
                  brightness_factor=1.2,
                  contrast_factor=1.2,
                  use_unsharp=True,
                  unsharp_radius=2,
                  unsharp_percent=150,
                  unsharp_threshold=3,
                  auto_contrast=True):
    """
    Enhances an image's clarity using default parameters.
    """
    # Open the image using PIL
    image = Image.open(io.BytesIO(image_bytes))
    
    # Convert to RGB if the image has an alpha channel to avoid issues with autocontrast
    if image.mode == 'RGBA':
        image = image.convert('RGB')
    
    # Apply enhancements in a specific order for best results
    # 1. First adjust brightness
    enhancer = ImageEnhance.Brightness(image)
    image = enhancer.enhance(brightness_factor)
    
    # 2. Then adjust contrast
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(contrast_factor)
    
    # 3. Apply auto contrast if requested
    if auto_contrast:
        try:
            image = ImageOps.autocontrast(image, cutoff=0.5)
        except Exception as e:
            print(f"Autocontrast failed: {e}, skipping this step")
    
    # 4. Sharpen the image
    enhancer = ImageEnhance.Sharpness(image)
    image = enhancer.enhance(sharpness_factor)
    
    # 5. Apply unsharp mask for additional sharpening if requested
    if use_unsharp:
        image = image.filter(ImageFilter.UnsharpMask(
            radius=unsharp_radius,
            percent=unsharp_percent,
            threshold=unsharp_threshold
        ))
    
    return image

def split_image_vertical(image: Image.Image) -> list:
    """
    Splits the image vertically into two halves.
    Returns a list containing the left and right halves.
    """
    width, height = image.size
    left_half = image.crop((0, 0, width // 2, height))
    right_half = image.crop((width // 2, 0, width, height))
    return [left_half, right_half]

def crop_top_portion(image: Image.Image, fraction=7/12) -> Image.Image:
    """
    Crops the top portion of an image based on the specified fraction.
    By default, crops to the top 7/12 of the image.
    Returns the cropped image.
    """
    width, height = image.size
    crop_height = int(height * fraction)
    cropped_image = image.crop((0, 0, width, crop_height))
    return cropped_image

def detect_image_type(image_bytes: bytes, api_key: str) -> str:
    """
    Detects the type of marketing mix model image.
    Returns 'one-pager' or 'budget-allocation' or 'unknown'.
    """
    try:
        # OCR based detection - look for key phrases
        image = Image.open(io.BytesIO(image_bytes))
        
        # Use pytesseract to extract text from the image
        text = pytesseract.image_to_string(image)
        
        # Print for debugging (will show in Streamlit logs)
        print(f"OCR text extracted: {text[:100]}...")
        
        # Print more of the OCR text for debugging
        print(f"Full OCR text: {text}")
        
        # First check for Budget Allocation (this should take precedence)
        if "Budget Allocation" in text or "Budget Allocation Onepager" in text or "Optimized Budget" in text:
            print("Detected as budget-allocation via OCR (primary check)")
            return "budget-allocation"
        
        # Then check for One-pager
        elif any(keyword in text for keyword in ["One-pager for Model", "Model Performance:", "Response Decomposition", "Response Curve", "NRMSE", "Adjusted R²"]):
            print("Detected as one-pager via OCR")
            return "one-pager"
        
        print("OCR detection failed, trying OpenAI Vision API")
        
        # If OCR doesn't find it, try the OpenAI Vision API for image detection
        prompt = """Classify this marketing mix model image. 

Look for these characteristics:

'One-pager for Model' typically includes:
- Model performance metrics (R²/NRMSE)
- Response curves or distribution charts
- Channel effectiveness/ROI measures
- Predictor importance graphs

'Budget Allocation Onepager' typically includes:
- Budget tables with spending amounts
- Channel allocation percentages
- Optimization scenarios
- ROAS comparisons across channels

Respond with ONLY ONE of these exact phrases:
- "one-pager" (for Model Performance/Response charts)
- "budget-allocation" (for Budget/Spending tables)"""
        
        client = OpenAI(api_key=api_key)
        base64_image = base64.b64encode(image_bytes).decode("utf-8")
        
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                temperature=0,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an image classification system that can only respond with 'one-pager' or 'budget-allocation'."
                    },
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
                            }
                        ]
                    }
                ],
                max_tokens=100
            )
            
            result = response.choices[0].message.content.lower().strip()
            print(f"OpenAI Vision API response: {result}")
            
            if "one-pager" in result or "one pager" in result:
                return "one-pager"
            elif "budget" in result or "allocation" in result:
                return "budget-allocation"
            else:
                print(f"Unrecognized API response format: {result}")
                return "unknown"
                
        except Exception as api_error:
            print(f"OpenAI API error in image detection: {str(api_error)}")
            # If API call fails, make best guess based on image characteristics
            
            # Check image size and aspect ratio as a fallback detection method
            width, height = image.size
            aspect_ratio = width / height
            
            # One-pagers tend to be wider (landscape orientation)
            if aspect_ratio > 1.2:
                print("Fallback: Using aspect ratio to guess image type (landscape: likely one-pager)")
                return "one-pager"
            else:
                print("Fallback: Using aspect ratio to guess image type (portrait/square: likely budget-allocation)")
                return "budget-allocation"
        
    except Exception as e:
        st.error(f"Error in image type detection: {str(e)}")
        # Last resort fallback
        return "unknown"

def pil_image_to_bytes(image: Image.Image, format="JPEG"):
    """
    Converts a Pillow Image to bytes.
    """
    img_byte_array = io.BytesIO()
    image.save(img_byte_array, format=format)
    return img_byte_array.getvalue()

def analyze_image_with_gpt4_vision_custom(image_bytes: bytes, custom_prompt: str, api_key: str) -> str:
    """
    Sends an image with a custom prompt to GPT-4 Vision for analysis.
    """
    try:
        client = OpenAI(api_key=api_key)
        base64_image = base64.b64encode(image_bytes).decode("utf-8")
        
        response = client.chat.completions.create(
            model="gpt-4o",
            temperature=0,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": custom_prompt},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
                        }
                    ]
                }
            ],
            max_tokens=1000
        )
        
        return response.choices[0].message.content
    except Exception as e:
        raise Exception(f"Error analyzing image with custom prompt: {str(e)}")

def summarize_responses(summary_prompt: str, api_key: str) -> str:
    """
    Uses GPT-4 to summarize combined responses.
    """
    try:
        client = OpenAI(api_key=api_key)
        
        response = client.chat.completions.create(
            model="gpt-4o",
            temperature=0,
            messages=[
                {
                    "role": "user",
                    "content": summary_prompt
                }
            ],
            max_tokens=1500
        )
        
        return response.choices[0].message.content
    except Exception as e:
        raise Exception(f"Error generating summary: {str(e)}")

def add_table_to_doc(doc: Document, table_lines: list):
    """
    Parses a list of strings (each representing a table row in markdown format)
    and adds a formatted table to the DOCX document.
    """
    if not table_lines:
        return
    
    # Parse the header row to determine column count
    header = table_lines[0]
    # Remove leading/trailing | and split by |
    columns = [col.strip() for col in header.strip('|').split('|')]
    col_count = len(columns)
    
    # Create table with appropriate number of rows and columns
    table = doc.add_table(rows=len(table_lines)-1, cols=col_count)
    table.style = 'Table Grid'
    
    # Process each row (skip the separator row)
    data_rows = [row for i, row in enumerate(table_lines) if i != 1]
    
    for i, row_text in enumerate(data_rows):
        # Remove leading/trailing | and split by |
        cells = [cell.strip() for cell in row_text.strip('|').split('|')]
        
        # Add data to table cells
        for j, cell_text in enumerate(cells):
            if j < col_count:  # Ensure we don't exceed column count
                cell = table.cell(i, j)
                cell.text = cell_text
                
                # Style header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(11)
    
    # Add some space after the table
    doc.add_paragraph()

def add_chart_to_doc(doc: Document, figure: Figure, has_negative_values: bool = False, title: str = "Response Contribution by Media Channel", negative_predictors: list = None):
    """
    Adds a matplotlib chart to the Word document.
    Saves the figure to a temporary file and then adds it to the document.
    If has_negative_values is True, adds a note about the presence of negative values.
    If negative_predictors is provided, lists the specific predictors with negative values.
    """
    # Create a heading for the chart
    doc.add_heading(title, level=3)
    
    # Save the figure to a temporary file
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
        figure.savefig(temp_file.name, format='png', bbox_inches='tight')
        # Add the image to the document
        doc.add_picture(temp_file.name, width=Inches(6.0))
    
    # Add explanation text
    p = doc.add_paragraph("This visualization shows the percentage contribution of each media channel to the overall response, based on the Response Decomposition Waterfall by Predictor from the one-pager model.")
    
    # Add note about negative values if needed
    if has_negative_values and negative_predictors:
        note_p = doc.add_paragraph()
        negative_predictors_str = ", ".join(negative_predictors)
        note_run = note_p.add_run(f"Note: Some predictors ({negative_predictors_str}) show negative contributions. This is common in marketing mix models where certain components may have negative contributions, indicating potential inefficiencies or cannibalization effects.")
        note_run.italic = True
    
    # Add a small space after the chart
    doc.add_paragraph()

def add_formatted_text_to_doc(doc: Document, text: str):
    """
    Processes the final_output string line by line.
    Lines that start with '###' are added as headings.
    Lines that look like table rows (starting and ending with '|') are collected and added as a table.
    All other lines are added as normal paragraphs.
    """
    # Split text into lines
    lines = text.strip().split('\n')
    
    # Track table lines
    table_lines = []
    collecting_table = False
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines
        if not line:
            # If we were collecting a table, add it now
            if collecting_table and table_lines:
                add_table_to_doc(doc, table_lines)
                table_lines = []
                collecting_table = False
            continue
        
        # Check if it's a table row
        if line.startswith('|') and line.endswith('|'):
            collecting_table = True
            table_lines.append(line)
            continue
        else:
            # If we were collecting a table, add it now
            if collecting_table and table_lines:
                add_table_to_doc(doc, table_lines)
                table_lines = []
                collecting_table = False
        
        # Check if it's a heading (starts with '#')
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            p.style.font.size = Pt(14)
            p.style.font.bold = True
            continue
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            p.style.font.size = Pt(16)
            p.style.font.bold = True
            continue
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.style.font.size = Pt(18)
            p.style.font.bold = True
            continue
        
        # Add as normal paragraph
        p = doc.add_paragraph(line)
    
    # If there's a table at the end, add it
    if table_lines:
        add_table_to_doc(doc, table_lines)

def chatbot_with_image(image_bytes: bytes, user_question: str, api_key: str, second_image_bytes: bytes = None, comprehensive_report: str = None) -> str:
    """
    Allows users to ask specific questions about marketing mix model images.
    Can analyze a single image or both One-pager and Budget Allocation images together.
    
    Args:
        image_bytes: Primary image bytes.
        user_question: The question asked by the user.
        api_key: OpenAI API key.
        second_image_bytes: Optional secondary image bytes for combined analysis.
        comprehensive_report: Optional comprehensive analysis report from previous analysis
                             to ensure consistency between Chat and Comprehensive modes.
    
    Returns:
        The AI response to the user's question.
    """
    try:
        client = OpenAI(api_key=api_key)
        
        if not api_key:
            raise ValueError("OpenAI API key is required.")
        
        # Process the first image based on its type
        first_image = None
        first_image_type = None
        try:
            first_image_type = detect_image_type(image_bytes, api_key)
            first_image = Image.open(io.BytesIO(image_bytes))
            
            # Process image based on detected type
            if first_image_type == "one-pager":
                # Split one-pager images vertically like in comprehensive analysis
                image_halves = split_image_vertical(first_image)
                
                # Create content list with placeholder for prompt text
                content_list = [
                    {"type": "text", "text": ""}  # Placeholder for the prompt text
                ]
                
                # Add both halves of the one-pager
                for half in image_halves:
                    half_bytes = pil_image_to_bytes(half)
                    base64_half = base64.b64encode(half_bytes).decode("utf-8")
                    content_list.append({
                        "type": "image_url",
                        "image_url": {"url": f"data:image/jpeg;base64,{base64_half}"}
                    })
                
                # Skip the standard image addition later
                first_image = None  
            elif first_image_type == "budget-allocation":
                # For budget allocation images, crop to top 7/12 for better focus on key data
                first_image = crop_top_portion(first_image, fraction=7/12)
                
                # Create the content list with placeholder
                content_list = [
                    {"type": "text", "text": ""}  # Placeholder for the prompt text
                ]
            else:
                # Create the content list with placeholder for unknown image types
                content_list = [
                    {"type": "text", "text": ""}  # Placeholder for the prompt text
                ]
                
        except Exception as e:
            print(f"Error processing first image in chat: {str(e)}")
            first_image = Image.open(io.BytesIO(image_bytes))  # Use original if processing fails
            
            # Create the content list with placeholder
            content_list = [
                {"type": "text", "text": ""}  # Placeholder for the prompt text
            ]
        
        # Add the processed first image if it wasn't a split one-pager
        if first_image:
            first_image_bytes = pil_image_to_bytes(first_image)
            base64_first_image = base64.b64encode(first_image_bytes).decode("utf-8")
            content_list.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{base64_first_image}"}
            })
        
        # If a second image is provided, process and add it too
        if second_image_bytes:
            second_image = None
            second_image_type = None
            try:
                second_image_type = detect_image_type(second_image_bytes, api_key)
                second_image = Image.open(io.BytesIO(second_image_bytes))
                
                # Process based on type
                if second_image_type == "one-pager":
                    # Split one-pager images vertically like in comprehensive analysis
                    image_halves = split_image_vertical(second_image)
                    
                    # Add both halves of the one-pager
                    for half in image_halves:
                        half_bytes = pil_image_to_bytes(half)
                        base64_half = base64.b64encode(half_bytes).decode("utf-8")
                        content_list.append({
                            "type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{base64_half}"}
                        })
                    
                    # Skip the standard image addition
                    second_image = None
                elif second_image_type == "budget-allocation":
                    second_image = crop_top_portion(second_image, fraction=7/12)
                    
            except Exception as e:
                print(f"Error processing second image in chat: {str(e)}")
                second_image = Image.open(io.BytesIO(second_image_bytes))  # Use original if processing fails
            
            # Add the processed second image if it wasn't a split one-pager
            if second_image:
                second_image_bytes_processed = pil_image_to_bytes(second_image)
                base64_second_image = base64.b64encode(second_image_bytes_processed).decode("utf-8")
                content_list.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{base64_second_image}"}
                })
        
        # Create appropriate prompt based on whether we have one or two images
        if second_image_bytes:
            prompt = f"""You have been provided with TWO marketing mix model analysis images:
1. A One-pager for Model - Contains model performance metrics like Adjusted R², NRMSE, media channel impacts, and ROI analysis.
2. A Budget Allocation Onepager - Shows budget allocation scenarios and optimizations.

The user asks: "{user_question}"

If the question is specifically about marketing mix models, media channels, ROAS, budgeting, or anything that might appear in these marketing images, please analyze BOTH images together to provide a comprehensive answer.
Look for relevant information in both images that could help answer the question more thoroughly.
If information from one image complements or contradicts information from the other, explain these connections.

If the question is about general knowledge, such as "What is machine learning?" or other topics not related to the marketing mix model analysis, you can use your general knowledge to provide a helpful, informative response without needing to reference the images.

When analyzing model performance:
- Carefully report the exact values for model quality metrics like Adjusted R² (look for this in the one-pager)
- Assess if the model is reliable based on these metrics (an Adjusted R² above 0.7 generally indicates a reliable model)
- Look for NRMSE (Normalized Root Mean Square Error) and DECOMP.RSSD values 

When analyzing channel performance and ROI:
- Extract specific ROAS values for each channel
- Pay close attention to channels with ROAS of 0 - these indicate channels with no return and should NOT be recommended for increased investment
- Channels with a ROAS of 0 should be recommended for re-evaluation or reduction in investment, not for increase
- Do NOT automatically recommend increasing investment based solely on high ROAS - check the ACTUAL scenarios in the budget allocation image first
- IMPORTANT: If the optimized scenarios in the budget allocation image show a channel receiving LESS investment in the bound_2 and bound_3 scenarios, you should recommend "Maintain" not "Increase", regardless of ROAS
- Note channels with the highest and lowest ROI
- Always check the Budget Allocation image's optimized scenarios to determine the final recommendation

For budget allocation recommendations:
- Be precise and data-driven - your recommendations must match what the data actually shows
- If a channel has 0 ROAS or 0 Expected Response, do not recommend increasing investment in it
- SPECIFIC DIRECTION FOR OOH CHANNEL: If OOH shows decreasing investment in the optimized budget scenarios, recommend "Maintain" not "Increase" 
- If the Budget Allocation data shows a channel getting reduced investment in optimized scenarios, do not recommend increasing investment even if ROAS is high
- Channels with 0 ROAS should be recommended for re-evaluation or reduction

Focus on providing actionable insights, specific metrics, and clear explanations that directly address what the user asked about.
Pay special attention to all numerical values, percentages, and statistical metrics visible in charts, tables, and annotations.
"""
        else:
            prompt = f"""The user has uploaded a marketing mix model analysis image and asks the following question:

"{user_question}"

If the question is specifically about marketing mix models, media channels, ROAS, budgeting, or anything that might appear in this marketing image, please analyze the image to provide a comprehensive answer.

If the question is about general knowledge, such as "What is machine learning?" or other topics not related to the marketing mix model analysis, you can use your general knowledge to provide a helpful, informative response without needing to reference the image.

When analyzing the marketing image:
- Focus on providing actionable insights, specific metrics, and clear explanations that directly address what the user asked about
- If the question is about model performance, make sure to extract and report specific model quality metrics like Adjusted R², NRMSE, and DECOMP.RSSD
- If the question is about ROI, ROAS, channel performance, budget allocation, or other model metrics, extract specific numbers and provide context
- Pay special attention to all numerical values, percentages, and statistical metrics visible in charts, tables, and annotations
- If the image doesn't contain information needed to answer the question completely, state what's missing and provide the best possible answer based on what's visible
"""
        
        # Update the first element (prompt text)
        content_list[0]["text"] = prompt
        
        # If we have a comprehensive report from previous analysis, include it
        if comprehensive_report:
            # Add information about the comprehensive report in the prompt
            comprehensive_report_prompt = f"""
I have already performed a comprehensive analysis of the marketing mix model images. 
Here is the detailed report I generated, which you should refer to when answering marketing-related questions:

COMPREHENSIVE ANALYSIS REPORT:
{comprehensive_report}

If the user's question is specifically about marketing mix models, media channels, ROAS, budgeting, or anything that appears in this analysis, please use this analysis as a foundation for your answer to ensure consistency with previous insights.

If the question is about general knowledge, such as "What is machine learning?" or other topics not related to the marketing mix model analysis, you can use your general knowledge to provide a helpful, informative response.

The user's question is: "{user_question}"
"""
            # Create a new message list with the comprehensive report first
            messages = [
                {"role": "system", "content": comprehensive_report_prompt},
                {"role": "user", "content": content_list}
            ]
        else:
            # If no comprehensive report, just use the content list
            messages = [
                {"role": "system", "content": "You are an AI assistant that specializes in marketing mix models, but can also help with general knowledge questions. When asked about marketing mix models, media channels, or marketing analytics, provide expert insights. For general knowledge questions, provide helpful and accurate information."},
                {"role": "user", "content": content_list}
            ]
        
        # Send to the API
        response = client.chat.completions.create(
            model="gpt-4o",
            temperature=0.0,
            messages=messages,
            max_tokens=1500
        )
        return response.choices[0].message.content
    except Exception as e:
        raise Exception(f"Error processing chat question: {str(e)}")

def analyze_one_pager_with_gpt4_vision(image_bytes: bytes, api_key: str) -> str:
    """
    Analyzes a One-pager for Model image with GPT-4 Vision.
    This is optimized for analyzing one half of the one-pager image.
    """
    try:
        # First, split the image vertically since one-pagers are dense with information
        image = Image.open(io.BytesIO(image_bytes))
        halves = split_image_vertical(image)
        
        # Create prompts for each half
        left_prompt = """You are an expert marketing data scientist analyzing a marketing mix model one-pager. 
        This is the LEFT HALF of a 'One-pager for Model' document.
        Analyze the charts and metrics shown in this half. Pay special attention to:
        
        1. Overall model metrics like Adjusted R², NRMSE, DECOMP.RSSD
           - IMPORTANT: If the Adjusted R² value is below 0.7, explicitly note that the model has limited explanatory power and may not be reliable for making significant business decisions. The lower the Adjusted R², the more cautious executives should be about implementing the recommendations.
        2. The Fitted vs. Residual plot
        3. Any response decomposition charts
        4. ROI/ROAS metrics for different channels
        
        Extract specific numbers, percentages, and metrics. Be precise in your analysis.
        Organize your analysis by metric/chart but don't list every single data point - focus on the important insights.
        """
        
        right_prompt = """You are an expert marketing data scientist analyzing a marketing mix model one-pager. 
        This is the RIGHT HALF of a 'One-pager for Model' document.
        Analyze the charts and metrics shown in this half. Pay special attention to:
        
        1. Adstock and carryover metrics (like Weibull PDF plots)
        2. Response curve analysis
        3. Any channel-specific metrics like ROAS/ROI not covered in the left half
        4. Any other charts or tables visible
        
        Extract specific numbers, percentages, and metrics. Be precise in your analysis.
        Organize your analysis by metric/chart but don't list every single data point - focus on the important insights.
        """
        
        # Process each half
        left_bytes = pil_image_to_bytes(halves[0])
        right_bytes = pil_image_to_bytes(halves[1])
        
        left_analysis = analyze_image_with_gpt4_vision_custom(left_bytes, left_prompt, api_key)
        right_analysis = analyze_image_with_gpt4_vision_custom(right_bytes, right_prompt, api_key)
        
        # Now create a comprehensive analysis combining both halves
        summary_prompt = f"""You are creating a comprehensive analysis of a marketing mix model one-pager.
        Below are two analyses: one of the left half and one of the right half of the same document.
        
        LEFT HALF ANALYSIS:
        {left_analysis}
        
        RIGHT HALF ANALYSIS:
        {right_analysis}
        
        Combine these analyses into a single, well-structured report that covers:
        
        ### Model Performance Overview
        Summarize the model's statistical quality (R², NRMSE, etc.)
        IMPORTANT: If the Adjusted R² value is below 0.7, explicitly note that the model has limited explanatory power and may not be reliable for making significant business decisions. The lower the Adjusted R², the more cautious executives should be about implementing the recommendations.
        
        ### Key Business Drivers
        Identify which channels are driving the most positive responses
        
        ### Investment Timing Impact
        Analyze the immediate vs. carryover effects
        
        ### ROI Analysis
        Compare ROI/ROAS across channels
        
        ### Optimization Recommendations
        Provide actionable suggestions based on the data
        
        Be specific with numbers, percentages, and insights. Maintain a structured format with clear headings.
        """
        
        final_analysis = summarize_responses(summary_prompt, api_key)
        return final_analysis
        
    except Exception as e:
        raise Exception(f"Error analyzing one-pager: {str(e)}")

def extract_predictor_percentages(image_bytes: bytes, api_key: str) -> dict:
    """
    Extracts predictor percentages from a Response Decomposition Waterfall chart.
    Returns a dictionary of predictor names and their percentage contributions.
    """
    try:
        # First, try to find the right section of the image that contains the waterfall chart
        # Split the image vertically since the chart is typically on the left half
        image = Image.open(io.BytesIO(image_bytes))
        halves = split_image_vertical(image)
        left_half_bytes = pil_image_to_bytes(halves[0])
        
        client = OpenAI(api_key=api_key)
        base64_image = base64.b64encode(left_half_bytes).decode("utf-8")
        
        prompt = """You are examining a marketing mix model one-pager, focusing on the Response Decomposition Waterfall by Predictor chart.
        
        This chart shows the contribution of each marketing channel/predictor to the total response.
        
        Please look at the Response Decomposition Waterfall by Predictor chart (if present) and extract:
        1. The exact name of each marketing channel/predictor (e.g., TV, Digital, Print, Radio)
        2. The exact percentage contribution of each channel (precise numbers as shown in the chart)
        3. IMPORTANT: Be sure to include the "intercept" value if it appears in the chart
        
        Format your response as a JSON object where keys are predictor names (exactly as labeled in the chart) 
        and values are their percentage contributions as decimal numbers (not strings with % symbols).
        
        Example format:
        {
          "intercept": 30.5,
          "TV": 28.4,
          "Digital": 15.2,
          "Print": 5.6
        }
        
        IMPORTANT: Only include actual predictors from the image with their exact percentages as shown.
        Include the "intercept" value that represents the baseline contribution.
        Do NOT round numbers or include values not explicitly shown. Do NOT make up data or include
        channels not present in the chart. If you can't find the chart or the percentages, return an empty object.
        """
        
        # Create a completion with the json_object response format
        response = client.chat.completions.create(
            model="gpt-4o",
            temperature=0,
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
                        }
                    ]
                }
            ],
            max_tokens=500
        )
        
        # Parse the JSON response
        result = json.loads(response.choices[0].message.content)
        
        # If no data was found, try again with the full image
        if not result:
            base64_full_image = base64.b64encode(image_bytes).decode("utf-8")
            second_response = client.chat.completions.create(
                model="gpt-4o",
                temperature=0,
                response_format={"type": "json_object"},
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/jpeg;base64,{base64_full_image}"}
                            }
                        ]
                    }
                ],
                max_tokens=500
            )
            result = json.loads(second_response.choices[0].message.content)
        
        # Log the extracted values for debugging
        print(f"Extracted predictor data: {result}")
        
        # We'll handle both positive and negative values in the create_predictor_chart function
        return result
        
    except Exception as e:
        print(f"Error extracting predictor percentages: {str(e)}")
        return {}

def get_filtered_values_note(has_negative_values: bool) -> str:
    """
    Generates the HTML note about negative values for display in the UI.
    If negative_predictors are available in session state, includes them in the note.
    """
    if not has_negative_values:
        return ""
        
    if 'negative_predictors' in st.session_state and st.session_state.negative_predictors:
        negative_predictors_str = ", ".join(st.session_state.negative_predictors)
        return f"""
        <p><b>Note:</b> Some predictors ({negative_predictors_str}) show negative contributions. 
        This is common in marketing mix models where certain components may have negative contributions, 
        indicating potential inefficiencies or cannibalization effects.</p>
        """
    else:
        return """
        <p><b>Note:</b> Some predictors show negative contributions. 
        This is common in marketing mix models where certain components may have negative contributions, 
        indicating potential inefficiencies or cannibalization effects.</p>
        """

def create_predictor_chart(predictor_data: dict) -> tuple:
    """
    Creates a horizontal bar chart visualization of predictor percentages.
    Returns a tuple containing (matplotlib Figure object, has_negative_values, negative_predictors).
    has_negative_values indicates if any negative values were found.
    negative_predictors is a list of predictors with negative values.
    
    Organizes the bars with positive values at the top and negative values at the bottom.
    """
    if not predictor_data:
        # Return an empty figure if no data
        fig = Figure(figsize=(10, 6))
        ax = fig.add_subplot(111)
        ax.text(0.5, 0.5, "No predictor data available", 
                horizontalalignment='center', verticalalignment='center', fontsize=14)
        ax.axis('off')
        return fig, False, []
    
    # Check if there are any negative values
    negative_predictors = [k for k, v in predictor_data.items() if v < 0]
    has_negative_values = len(negative_predictors) > 0
    
    # Separate positive and negative values
    positive_items = {k: v for k, v in predictor_data.items() if v >= 0}
    negative_items = {k: v for k, v in predictor_data.items() if v < 0}
    
    # Sort positive values by value (ascending - to put highest values at top when plotted)
    # This counterintuitive ordering is because matplotlib plots from bottom to top
    sorted_positive = dict(sorted(positive_items.items(), key=lambda item: item[1]))
    
    # Sort negative values by value (descending - to put most negative at bottom)
    sorted_negative = dict(sorted(negative_items.items(), key=lambda item: item[1], reverse=True))
    
    # Combine negative first (will be at bottom), then positive values (will be at top)
    # This ordering ensures the highest positive values appear at the top of the plot
    combined_data = {**sorted_negative, **sorted_positive}
    
    # Prepare data for the bar chart
    labels = list(combined_data.keys())
    values = list(combined_data.values())
    
    # Create the bar chart
    fig = Figure(figsize=(12, 8))  # Larger figure to accommodate horizontal bars and labels
    ax = fig.add_subplot(111)
    
    # Custom colors with good contrast
    colors = plt.cm.tab20.colors
    
    # Assign colors to bars based on positive/negative values
    bar_colors = [colors[i % len(colors)] if v >= 0 else 'r' for i, v in enumerate(values)]
    
    # Create horizontal bar chart
    bars = ax.barh(labels, values, color=bar_colors, edgecolor='w', linewidth=0.5)
    
    # Add value labels at the end of each bar
    for i, bar in enumerate(bars):
        width = bar.get_width()
        label_x_pos = width + 0.5 if width >= 0 else width - 1.5
        ax.text(label_x_pos, bar.get_y() + bar.get_height()/2, 
                f'{values[i]:.1f}%', va='center', ha='left' if width >= 0 else 'right',
                fontweight='bold', color='black')
    
    # Add grid lines for better readability
    ax.grid(True, axis='x', linestyle='--', alpha=0.7)
    
    # Set title and labels
    ax.set_title('Response Contribution by Media Channel', fontsize=16)
    ax.set_xlabel('Contribution (%)', fontsize=12)
    
    # Add zero line for reference with negative values
    if has_negative_values:
        ax.axvline(x=0, color='black', linestyle='-', alpha=0.5)
    
    # Adjust layout
    fig.tight_layout()
    
    # Return the figure, a boolean indicating if negative values were found, and list of negative predictors
    return (fig, has_negative_values, negative_predictors)

def analyze_budget_allocation_with_gpt4_vision(image_bytes: bytes, api_key: str) -> str:
    """
    Analyzes a Budget Allocation Onepager for Model image with GPT-4 Vision.
    This is optimized specifically for budget allocation data.
    """
    try:
        # Create a focused prompt specifically for budget allocation data
        prompt = """You are an expert marketing analyst reviewing a 'Budget Allocation Onepager for Model' document.
        
        Analyze this image in detail and extract the following key information:
        
        1. Budget Allocation Tables:
           - Extract the complete 'Budget Allocation per Paid Media Variable per Month' table data
           - Create a comparison table showing channel performance across different channels:
           
           | Paid Media | abs.mean spend ($) | mean spend% | mean response% | mean ROAS | mROAS |
           |------------|-------------------|-------------|----------------|-----------|-------|
           | [Channel 1]| [Value]           | [Value]     | [Value]        | [Value]   | [Value] |
           | [Channel 2]| [Value]           | [Value]     | [Value]        | [Value]   | [Value] |
           ... and so on for all channels
        
        2. Budget Optimization Scenarios:
           
           ### i) Maintain Current Spend Scenario
           
           Present ONLY a concise table of the current spending allocation:
           
           | Channel | Current Spend ($) | Current Spend (%) | ROAS |
           |---------|-------------------|-------------------|------|
           | Channel 1 | [Value]         | [Value]           | [Value] |
           
           ### ii) Optimization Recommendations
           
           Instead of detailed tables for the 10% increase and 10% decrease scenarios, provide ONLY bullet-point recommendations for which channels should receive more or less investment:
           
           **Channels to Increase Investment:**
           * [Channel Name] - Brief justification based on ROAS
           * [Channel Name] - Brief justification based on ROAS
           
           **Channels to Decrease Investment:**
           * [Channel Name] - Brief justification based on ROAS
           * [Channel Name] - Brief justification based on ROAS
           
           ### iii) Final Recommendation Table
           
           Provide a final recommendation table showing the optimal strategy for each channel:
           
           | Channel | Recommendation |
           |---------|----------------|
           | Channel 1 | [Increase/Decrease/Maintain/Re-evaluate] |
        
        Be extremely specific with numbers and percentages. Create a structured analysis with clear, actionable insights.
        Ensure all tables align perfectly and contain accurate data extracted from the image.
        """
        
        # Send the full image for analysis
        analysis = analyze_image_with_gpt4_vision_custom(image_bytes, prompt, api_key)
        
        # Structure the output
        structured_output = f"""## Budget Allocation Analysis
        
        {analysis}
        
        """
        
        return structured_output
        
    except Exception as e:
        raise Exception(f"Error analyzing budget allocation: {str(e)}")

def comprehensive_analysis(one_pager_image: Image.Image, budget_image: Image.Image, api_key: str) -> tuple:
    """
    Performs a comprehensive analysis of both the One-pager for Model and 
    Budget Allocation Onepager images. 
    
    Processes each image appropriately (split vertically for one-pager,
    crop to 7/12 for budget allocation) and then generates a comprehensive
    report covering:
    
    1. Model Performance overview
    2. Key business drivers
    3. Investment Timing Impact
    4. ROI for each media channel
    5. Budget Allocation optimization scenarios
    6. Summary and recommendations
    
    Returns a tuple containing (final_report, predictor_data, chart_figure, has_negative_values, negative_predictors)
    """
    try:
        # Process and analyze the one-pager
        one_pager_bytes = pil_image_to_bytes(one_pager_image)
        one_pager_analysis = analyze_one_pager_with_gpt4_vision(one_pager_bytes, api_key)
        
        # Extract predictor percentages for visualization
        predictor_data = extract_predictor_percentages(one_pager_bytes, api_key)
        chart_result = create_predictor_chart(predictor_data)
        # Unpack the tuple (fig, has_negative_values, negative_predictors)
        pie_chart, has_negative_values, negative_predictors = chart_result
        
        # Save in session state for UI display
        st.session_state.negative_predictors = negative_predictors
        
        # Process and analyze the budget allocation image
        # First crop to top 7/12 as this contains the most relevant budget information
        cropped_budget_image = crop_top_portion(budget_image, fraction=7/12)
        budget_bytes = pil_image_to_bytes(cropped_budget_image)
        budget_analysis = analyze_budget_allocation_with_gpt4_vision(budget_bytes, api_key)
        
        # Combine the analyses
        summary_prompt = f"""You are preparing a comprehensive marketing mix model report for executives.
        You have two sections of analysis below - one from the model performance one-pager and one from the budget allocation one-pager.
        
        MODEL PERFORMANCE ANALYSIS:
        {one_pager_analysis}
        
        BUDGET ALLOCATION ANALYSIS:
        {budget_analysis}
        
        Create a cohesive, executive-ready report that combines these analyses with the following structure:
        
        # Marketing Mix Model AI-Powered Analysis
        
        ## 1. Model Performance Overview
        Summarize the key model quality metrics and what they tell us about the reliability of the model.
        
        IMPORTANT: If the Adjusted R² value is below 0.7, explicitly note that the model has limited explanatory power and may not be reliable for making significant business decisions. The lower the Adjusted R², the more cautious executives should be about implementing the recommendations.
        
        ## 2. Key Business Drivers
        Identify which marketing channels and other factors are driving business results.
        
        ## 3. Investment Timing Impact
        Explain the immediate vs. carryover effects of marketing investments and implications for campaign timing.
        
        ## 4. Channel ROI Analysis
        Compare ROI/ROAS across all channels with specific numbers.
        Pay close attention to channels with ROAS of 0 - these indicate channels with no return and should NOT be recommended for increased investment.
        Channels with a ROAS of 0 should be recommended for re-evaluation or reduction in investment, not for increase.
        Only recommend increasing investment in channels with proven positive ROAS.
        
        ## 5. Budget Optimization Scenarios
        
        ### 5.1. Maintain Current Spend Scenario
        Present ONLY a concise table of the current spending allocation:
        
        | Channel | Current Spend ($) | Current Spend (%) | ROAS |
        |---------|-------------------|-------------------|------|
        | Channel 1 | [Value]         | [Value]           | [Value] |
        
        ### 5.2. Optimization Recommendations
        
        Instead of detailed tables for the 10% increase and 10% decrease scenarios, provide ONLY bullet-point recommendations for which channels should receive more or less investment. These recommendations must match the final recommendation table in section 5.3:
        
        **Channels to Increase Investment:**
        * [Channel Name] - Brief justification based on both ROAS and the bounded scenario results
        * [Channel Name] - Brief justification based on both ROAS and the bounded scenario results
        
        **Channels to Decrease Investment or Re-evaluate:**
        * [Channel Name] - Brief justification based on both ROAS and the bounded scenario results
        * [Channel Name] - Brief justification based on both ROAS and the bounded scenario results
        
        **Channels to Maintain:**
        * [Channel Name] - Brief justification based on both ROAS and the bounded scenario results
        
        ### 5.3. Final Recommendation Table
        
        Provide a final recommendation table showing the optimal strategy for each channel:
        
        | Channel | Recommendation |
        |---------|----------------|
        | Channel 1 | [Increase/Decrease/Maintain/Re-evaluate] |
        
        CRITICAL WHEN MAKING RECOMMENDATIONS:
        1. Look at the "Budget Allocation per Paid Media Variable per Month" table carefully.
        2. If a channel shows an increase in the "Bounded" scenario compared to "Initial", this indicates it should receive MORE investment.
        3. If a channel shows a decrease in the "Bounded" scenario compared to "Initial", this indicates it should receive LESS investment.
        4. Base your recommendations primarily on what the optimization actually shows in the table, not just on ROAS values.
        
        For example, if edu.spend increases from $12K to $13K in the bounded scenario, it should be recommended for increase.
        If web.spend decreases from $1.9K to $1.7K in the bounded scenario, it should not be recommended for increase.
        
        IMPORTANT: Look carefully at the OOH_s (Out-of-Home) channel specifically. If OOH_s spend decreases in the "Bounded" scenario compared to "Initial", it should NOT be recommended for increase, even if it has a positive ROAS. The optimization model is indicating that reducing OOH_s investment would be more optimal.
        
        The Final Recommendation table should clearly indicate:
        - Channels with 0 ROAS should be marked as "Re-evaluate" 
        - Channels showing increased spend in optimized scenarios should be marked as "Increase"
        - Channels showing decreased spend in optimized scenarios should be marked as "Decrease"
        - Channels with unchanged spend should be marked as "Maintain"
        
        ## 6. Summary and Strategic Recommendations
        Provide 3-5 clear, actionable recommendations based on both analyses.
        
        IMPORTANT FOR RECOMMENDATIONS:
        - Never recommend increasing investment in channels with 0 ROAS
        - For channels with 0 ROAS (like Facebook if shown with 0), recommend re-evaluation, not increased investment
        - Only recommend increasing investment in channels that show increased spend in the optimized scenarios
        - HIGHEST PRIORITY FOR TV_S: For tv_s (Television) specifically, ALWAYS verify if there is ANY increase from Initial to Bounded scenario, even a small one. If tv_s shows ANY increase at all OR has a ROAS above 2.0, it MUST be recommended as "Increase" in BOTH the bullet points and final table due to its high effectiveness and strategic importance. This rule overrides all other rules for the tv_s channel specifically.
        - CRITICAL: For OOH_s (Out-of-Home) specifically, verify the optimized scenarios very carefully. If OOH_s shows ANY decrease or is not explicitly increased in the optimized scenarios, it MUST be labeled as "Re-evaluate" or "Decrease", NOT as "Increase". This is a crucial check for accuracy.
        - Be precise in your channel names and make sure recommendations match the data in tables
        - Double-check your recommendations against the budget allocation tables to ensure consistency
        
        Be executive-ready: concise but comprehensive, with specific numbers and actionable insights.
        Format any tables cleanly and ensure they add clarity to your insights.
        """
        
        final_report = summarize_responses(summary_prompt, api_key)
        # Return the report, predictor data, pie chart figure, has_negative_values flag, and negative_predictors
        return (final_report, predictor_data, pie_chart, has_negative_values, negative_predictors)
            
    except Exception as e:
        raise Exception(f"Error generating comprehensive analysis: {str(e)}")

def analyze_image_with_gpt4_vision(image_bytes: bytes, api_key: str) -> str:
    """
    General analyze function that detects the image type and uses the appropriate specialized analyzer.
    This is maintained for compatibility with the chat mode and any other parts of the code that expect this function.
    """
    # First, try to detect the image type and use specialized analyzers
    try:
        image_type = detect_image_type(image_bytes, api_key)
        
        if image_type == "one-pager":
            return analyze_one_pager_with_gpt4_vision(image_bytes, api_key)
        elif image_type == "budget-allocation":
            # Process the budget allocation image - crop to top 7/12 for better analysis
            image = Image.open(io.BytesIO(image_bytes))
            cropped_image = crop_top_portion(image, fraction=7/12)
            cropped_bytes = pil_image_to_bytes(cropped_image)
            return analyze_budget_allocation_with_gpt4_vision(cropped_bytes, api_key)
    except Exception as e:
        # Log the error but continue with generic analysis
        print(f"Error in image type detection: {str(e)}")
    
    # If we get here, either image type detection failed or it's an unknown type
    # Use generic analysis as a fallback
    detailed_prompt = """You are an experienced marketing data scientist. The user has uploaded an image that appears to be from a marketing mix model.
    Please analyze this image in detail, focusing on any metrics, charts, graphs, or tables visible.
    Extract all relevant information about marketing performance, ROI, channel efficiency, or budget allocation that you can see.
    Provide a comprehensive analysis of what this image shows about the marketing mix model.
    
    If the image is titled "One-pager for Model", interpret the key insights from the charts provided and address:
    - Overall Model Performance (metrics like Adjusted R², NRMSE)
      - IMPORTANT: If the Adjusted R² value is below 0.7, explicitly note that the model has limited explanatory power and may not be reliable for making significant business decisions
    - Channel Contributions and Efficiency 
    - Time Dynamics and Carryover Effects
    - Spend and Response Relationships
    - Actionable Recommendations
    
    If the image is titled "Budget Allocation Onepager", please extract:
    - Budget allocation data across different scenarios
    - Channel-specific metrics (spend, response, ROAS)
    - Optimization recommendations 
    - Present a focused budget analysis in the following format:
    
    ### i) Maintain Current Spend Scenario
           
    Present ONLY a concise table of the current spending allocation:
    
    | Channel | Current Spend ($) | Current Spend (%) | ROAS |
    |---------|-------------------|-------------------|------|
    | Channel 1 | [Value]         | [Value]           | [Value] |
    
    ### ii) Optimization Recommendations
    
    Instead of detailed tables for the 10% increase and 10% decrease scenarios, provide ONLY bullet-point recommendations for which channels should receive more or less investment. These recommendations must match the final recommendation table in section 5.3:
    
    **Channels to Increase Investment:**
    * [Channel Name] - Brief justification based on both ROAS and the bounded scenario results
    * [Channel Name] - Brief justification based on both ROAS and the bounded scenario results
    
    **Channels to Decrease Investment or Re-evaluate:**
    * [Channel Name] - Brief justification based on both ROAS and the bounded scenario results
    * [Channel Name] - Brief justification based on both ROAS and the bounded scenario results
    
    **Channels to Maintain:**
    * [Channel Name] - Brief justification based on both ROAS and the bounded scenario results
    
    ### iii) Final Recommendation Table
    
    Provide a final recommendation table showing the optimal strategy for each channel:
    
    | Channel | Recommendation |
    |---------|----------------|
    | Channel 1 | [Increase/Decrease/Maintain/Re-evaluate] |
    
    3. Final Recommendations:
       After analyzing all three scenarios, provide a final recommendation table:
       
       | Channel | Optimal Strategy |
       |---------|------------------|
       | Channel 1 | [Strategy] |
       
       Follow these rules when making recommendations:
       - Channels with 0 ROAS should ALWAYS be marked as "Re-evaluate"
       - HIGHEST PRIORITY FOR TV_S: For tv_s (Television) specifically, ALWAYS verify if there is ANY increase from Initial to Bounded scenario, even a small one. If tv_s shows ANY increase at all OR has a ROAS above 2.0, it MUST be recommended as "Increase" in BOTH the bullet points and final table due to its high effectiveness and strategic importance. This rule overrides all other rules for the tv_s channel specifically.
       - Pay extremely close attention to any OOH_s channel (Out-of-Home) - NEVER recommend "Increase" for OOH_s unless it explicitly shows increased investment in the optimized scenarios
       - If OOH_s (Out-of-Home) shows ANY decrease in spend in the optimized scenarios or receives less allocation, it MUST be marked as "Decrease" or "Re-evaluate", NEVER as "Increase"
       - Pay close attention to Search channel - if it shows 0 in the optimized budget scenarios, recommend "Re-evaluate" not "Increase" for Search
       - Only recommend "Increase" for channels where the data clearly shows positive ROI AND increased investment in optimized scenarios
       - For Facebook, if it shows 0 ROAS, it must be marked as "Re-evaluate"
       - For Search, if it shows 0 in the budget allocation scenarios, it must be marked as "Re-evaluate" not "Increase"
       - If a channel has high ROAS but gets LESS investment or 0 investment in the optimized scenarios, recommend "Re-evaluate" or "Maintain", never "Increase"
    """
    
    try:
        client = OpenAI(api_key=api_key)
        if not api_key:
            raise ValueError("OpenAI API key is required.")
            
        base64_image = base64.b64encode(image_bytes).decode("utf-8")
        response = client.chat.completions.create(
            model="gpt-4o",
            temperature=0,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": detailed_prompt},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
                        }
                    ]
                }
            ],
            max_tokens=1000
        )
        return response.choices[0].message.content
    except Exception as e:
        raise Exception(f"Error analyzing image: {str(e)}")

def main():
    # Initialize session state for API key if not exists
    if 'api_key' not in st.session_state:
        st.session_state.api_key = ''
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []

    # Load CSS styles
    if os.path.exists("styles.css"):
        with open("styles.css") as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

    # Add sidebar with instructions
    with st.sidebar:
        st.title("Welcome to the LLM-Powered AI Marketing Mix Analyzer Tool")
        st.write("This application helps you analyze marketing mix model outputs using advanced AI. Here's how to use it:")
        
        st.markdown("**1. Select Mode:** Choose your preferred analysis mode:")
        st.markdown("   • **AI-Powered Analysis:** For detailed analysis of both model images")
        st.markdown("   • **Chat with Model:** For interactive Q&A about your model")
        
        st.markdown("**2. Upload Image(s):** Upload marketing mix model output image(s)")
        st.markdown("**3. Enter API Key:** Provide your OpenAI API key (securely masked)")
        st.markdown("**4. Get Results:** Receive detailed insights and recommendations")
        
        st.markdown("For optimal results, use high-resolution images showing clear charts and metrics.")
        
        st.subheader("AI-Powered Analysis Details")
        st.markdown("This mode requires two specific images:")
        st.markdown("• **One-pager for Model:** Contains performance metrics, ROI, and response charts")
        st.markdown("• **Budget Allocation Onepager:** Contains budget allocation scenarios")
        
        st.markdown("The output will include all six key sections: Model Performance, Business Drivers, Timing Impact, ROI Analysis, Budget Scenarios, and Recommendations.")
    
    st.title("🤖 LLM-Powered AI Marketing Mix Analyzer")

    # Initialize the mode selection in session state if not already set
    if 'mode_selection' not in st.session_state:
        st.session_state.mode_selection = "AI-Powered Analysis"
    
    # Check if we need to switch to AI-Powered Analysis mode
    if 'switch_to_comprehensive' in st.session_state and st.session_state.switch_to_comprehensive:
        # Remove the flag
        del st.session_state.switch_to_comprehensive
        # Force the index to be 0 (AI-Powered Analysis)
        mode_index = 0
    else:
        # Use current selection
        mode_index = 0 if st.session_state.mode_selection == "AI-Powered Analysis" else 1
    
    # Mode selection with session state
    analysis_mode = st.selectbox(
        "Select Analysis Mode",
        ["AI-Powered Analysis", "Chat with Model"],
        index=mode_index,
        help="Choose between comprehensive analysis with two images or interactive Q&A",
        key="mode_selection"
    )
    
    # Get API key from environment variables - define this early so it's available in all modes
    env_api_key = os.environ.get("OPENAI_API_KEY", "")
    
    # Add a note about the API key
    if env_api_key:
        st.info("An API key is available from environment variables, but you need to enter your own key below.")
    
    # Always show a masked placeholder for the API key
    placeholder = "XXXXXXXXXXXXXXXXXXXXXXXXXXXX"
    
    # Always show the API key input field with an empty value (never prefill)
    api_key = st.text_input(
        "Enter your OpenAI API Key",
        value="",
        placeholder=placeholder,
        type="password",
        help="Your API key will not be stored and needs to be entered each time you use the application.",
        key="api_key_input"
    )
    
    # If the environment variable exists but the user hasn't entered anything, use the environment variable
    if not api_key and env_api_key:
        api_key = env_api_key

    if analysis_mode == "AI-Powered Analysis":
        # For AI-powered analysis, we need both types of images
        st.markdown("""
        <div class="info-box">
            <h3>AI-Powered Analysis Mode</h3>
            <p>Please upload both required model output images:</p>
            <ol>
                <li><b>"One-pager for Model"</b> - Contains performance metrics, ROI, and response charts</li>
                <li><b>"Budget Allocation Onepager for Model"</b> - Contains budget allocation scenarios</li>
            </ol>
        </div>
        """, unsafe_allow_html=True)
        
        # Check if images are already in session state
        has_one_pager = 'one_pager_bytes' in st.session_state and st.session_state.one_pager_bytes is not None
        has_budget = 'budget_bytes' in st.session_state and st.session_state.budget_bytes is not None
        
        # Two separate upload fields for each type of image
        with st.container():
            col1, col2 = st.columns(2)
            
            with col1:
                one_pager_file = st.file_uploader(
                    "Upload One-pager for Model (PNG, JPG, JPEG)",
                    type=["png", "jpg", "jpeg"],
                    help="The One-pager contains model metrics, ROI, and response charts. Drag and drop your file here.",
                    key="one_pager_uploader",
                    label_visibility="visible"
                )
            
            with col2:
                budget_file = st.file_uploader(
                    "Upload Budget Allocation Onepager (PNG, JPG, JPEG)",
                    type=["png", "jpg", "jpeg"],
                    help="The Budget Allocation Onepager contains budget optimization scenarios. Drag and drop your file here.",
                    key="budget_uploader",
                    label_visibility="visible"
                )
        
        # Initialize variables for image processing
        one_pager_image = None
        one_pager_bytes = None
        budget_image = None
        budget_bytes = None
        new_images_uploaded = False
        
        # Check if we have new uploads and process them
        if one_pager_file or budget_file:
            # Clear previous report if new images are uploaded
            if 'report' in st.session_state:
                del st.session_state.report
                if 'predictor_data' in st.session_state:
                    del st.session_state.predictor_data
                if 'pie_chart' in st.session_state:
                    del st.session_state.pie_chart
        
        # Process one-pager image if uploaded
        if one_pager_file:
            with st.spinner("Processing One-pager for Model..."):
                # Read the image bytes
                one_pager_bytes = one_pager_file.read()
                one_pager_image = enhance_image(one_pager_bytes)
                
                # Store in session state
                st.session_state.one_pager_bytes = one_pager_bytes
                st.session_state.one_pager_image = one_pager_image
                st.session_state.one_pager_filename = one_pager_file.name
                
                # Display the one-pager image
                st.image(one_pager_image, use_container_width=True, caption="One-pager for Model")
                new_images_uploaded = True
        
        # Process budget allocation image if uploaded
        if budget_file:
            with st.spinner("Processing Budget Allocation Onepager..."):
                # Read the image bytes
                budget_bytes = budget_file.read()
                budget_image = enhance_image(budget_bytes)
                
                # Store in session state
                st.session_state.budget_bytes = budget_bytes
                st.session_state.budget_image = budget_image
                st.session_state.budget_filename = budget_file.name
                
                # Display the budget image
                st.image(budget_image, use_container_width=True, caption="Budget Allocation Onepager")
                new_images_uploaded = True
                
        # If no new images were uploaded, use the ones from session state if available
        if not new_images_uploaded:
            if has_one_pager:
                st.image(st.session_state.one_pager_image, use_container_width=True, caption="One-pager for Model (from previous upload)")
                one_pager_image = st.session_state.one_pager_image
                one_pager_bytes = st.session_state.one_pager_bytes
            
            if has_budget:
                st.image(st.session_state.budget_image, use_container_width=True, caption="Budget Allocation Onepager (from previous upload)")
                budget_image = st.session_state.budget_image
                budget_bytes = st.session_state.budget_bytes
        
        # Images are now stored in session state at upload time
        # The analysis_completed flag can be used if needed to indicate a full analysis has been run
            
        # Setup for Chat with Model mode
        uploaded_file = None
        enhanced_image = None
        
    else:  # Chat with Model mode
        # This will be handled by the new code in main()
        # Initialize default values
        uploaded_file = None
        enhanced_image = None

    # API key has already been defined above

    # Define image availability for both modes
    if analysis_mode == "AI-Powered Analysis":
        # Check for one_pager availability
        one_pager_available = (
            ('one_pager_image' in locals() and one_pager_image is not None) or
            ('one_pager_bytes' in st.session_state and st.session_state.one_pager_bytes is not None)
        )
        
        # Check for budget availability
        budget_available = (
            ('budget_image' in locals() and budget_image is not None) or
            ('budget_bytes' in st.session_state and st.session_state.budget_bytes is not None)
        )
    else:
        one_pager_available = False
        budget_available = False
        
    # Show warning if an image is uploaded but no API key is provided
    if not api_key and (
        (analysis_mode != "AI-Powered Analysis" and uploaded_file is not None) or 
        (analysis_mode == "AI-Powered Analysis" and (one_pager_available or budget_available))
    ):
        st.warning("Please enter your OpenAI API key to analyze the image(s).")
        st.markdown("""
        <div class="info-box">
            <p>To get your OpenAI API key:</p>
            <ol>
                <li>Go to <a href="https://platform.openai.com/account/api-keys" target="_blank">OpenAI API Keys</a></li>
                <li>Sign in or create an account</li>
                <li>Create a new API key</li>
                <li>Copy and paste the key here</li>
            </ol>
        </div>
        """, unsafe_allow_html=True)

    # Analysis options based on selected mode
    if analysis_mode == "AI-Powered Analysis":
        # AI-Powered Analysis mode requires both images to be uploaded
        st.markdown("""
        <div class="info-box">
            <h4>AI-Powered Analysis will provide:</h4>
            <ol>
                <li>Model Performance overview (Actual vs Predicted Response)</li>
                <li>Key business drivers / consumption contributors</li>
                <li>Investment Timing Impact (Immediate vs Carryover)</li>
                <li>ROI for each media channel</li>
                <li>Budget Allocation optimization Scenarios (Maintain Current Spend, if the business has extra 10% increase in budget which channels should receive more investment for better ROI, if the business cuts 10% in budget which channels should receive less investment while maintaining optimal ROI)</li>
                <li>Summary of the recommendations and insights</li>
            </ol>
        </div>
        """, unsafe_allow_html=True)
        
        # Update the availability checks for the comprehensive analysis button
        comprehensive_ready = (
            one_pager_available and
            budget_available and
            api_key
        )
        
        # Check if we have a previously generated report
        has_previous_report = (
            'report' in st.session_state and 
            'predictor_data' in st.session_state and 
            'pie_chart' in st.session_state and
            'has_negative_values' in st.session_state
        )
        
        # Show the cached report if available
        if has_previous_report:
            st.info("Showing previously generated report. Click 'Run AI-Powered Analysis' to regenerate if needed.")
            
            # Get the report and data from session state
            report = st.session_state.report
            predictor_data = st.session_state.predictor_data
            pie_chart = st.session_state.pie_chart
            has_negative_values = st.session_state.has_negative_values
            
            # Show the report
            st.markdown("## 📋 AI-Powered Analysis Report")
            
            # Split the report into sections to insert the pie chart after Model Performance section
            report_sections = report.split("## 2. Key Business Drivers")
            
            if len(report_sections) > 1:
                # Display first section (Model Performance)
                st.markdown(report_sections[0])
                
                # Display the pie chart after Model Performance section
                if predictor_data:
                    st.markdown("### 📈 Response Contribution by Media Channel")
                    st.pyplot(pie_chart)
                    
                    # Base info text
                    info_text = """
                    <div class="info-box">
                        <p>This visualization shows the percentage contribution of each media channel to the overall response, 
                        based on the Response Decomposition Waterfall by Predictor from the one-pager model.</p>
                    """
                    
                    # Add explanation about filtered negative values if needed
                    if has_negative_values:
                        # Use the helper function to get the note with specific predictor names
                        filtered_note = get_filtered_values_note(has_negative_values)
                        info_text += filtered_note
                    
                    info_text += "</div>"
                    st.markdown(info_text, unsafe_allow_html=True)
                
                # Display the rest of the report
                st.markdown("## 2. Key Business Drivers" + report_sections[1])
            else:
                # Fallback if report doesn't have expected structure
                st.markdown(report)
                
                # Display the pie chart if data is available
                if predictor_data:
                    st.markdown("### 📈 Response Contribution by Media Channel")
                    st.pyplot(pie_chart)
                    
                    # Base info text
                    info_text = """
                    <div class="info-box">
                        <p>This visualization shows the percentage contribution of each media channel to the overall response, 
                        based on the Response Decomposition Waterfall by Predictor from the one-pager model.</p>
                    """
                    
                    # Add explanation about filtered negative values if needed
                    if has_negative_values:
                        # Use the helper function to get the note with specific predictor names
                        filtered_note = get_filtered_values_note(has_negative_values)
                        info_text += filtered_note
                    
                    info_text += "</div>"
                    st.markdown(info_text, unsafe_allow_html=True)
                    
            # Option to download as DOCX
            doc = Document()
            
            # Split the report to insert the pie chart after Model Performance section
            report_sections = report.split("## 2. Key Business Drivers")
            
            if len(report_sections) > 1:
                # Add first section (Model Performance)
                add_formatted_text_to_doc(doc, report_sections[0])
                
                # Add the pie chart after Model Performance section
                if predictor_data:
                    add_chart_to_doc(doc, pie_chart, has_negative_values)
                
                # Add the rest of the report
                add_formatted_text_to_doc(doc, "## 2. Key Business Drivers" + report_sections[1])
            else:
                # Fallback if report doesn't have expected structure
                add_formatted_text_to_doc(doc, str(report))
                
                # Add the pie chart at the end if structure not recognized
                if predictor_data:
                    add_chart_to_doc(doc, pie_chart, has_negative_values)
            
            # Save temporarily and offer for download
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                doc.save(temp_file.name)
                
                with open(temp_file.name, 'rb') as file:
                    docx_bytes = file.read()
                    
                st.download_button(
                    label="📄 Download Report as DOCX",
                    data=docx_bytes,
                    file_name="Marketing_Mix_Model_Analysis.docx",
                    help="Download the analysis as a Word document",
                    key="download_cached_report"
                )
        
        # Always show the analyze button to allow regenerating the report
        analyze_button = st.button(
            "🔍 Run AI-Powered Analysis",
            help="Click to analyze both images and generate a comprehensive report",
            disabled=not comprehensive_ready
        )
        
        if analyze_button:
            with st.spinner("Analyzing marketing mix model images... This may take a minute or two."):
                try:
                    # Images are already stored in session state at upload time
                    # We'll keep this flag to indicate that analysis has been run
                    st.session_state.analysis_completed = True
                    
                    # Check if we already have a stored report
                    if ('report' in st.session_state and 
                        'predictor_data' in st.session_state and 
                        'pie_chart' in st.session_state and
                        'has_negative_values' in st.session_state):
                        # Use cached report data
                        report = st.session_state.report
                        predictor_data = st.session_state.predictor_data
                        pie_chart = st.session_state.pie_chart
                        has_negative_values = st.session_state.has_negative_values
                    else:
                        # Run the comprehensive analysis
                        report, predictor_data, pie_chart, has_negative_values, negative_predictors = comprehensive_analysis(
                            one_pager_image,
                            budget_image,
                            api_key
                        )
                        
                        # Store the report and data in session state
                        st.session_state.report = report
                        st.session_state.predictor_data = predictor_data
                        st.session_state.pie_chart = pie_chart
                        st.session_state.has_negative_values = has_negative_values
                        st.session_state.negative_predictors = negative_predictors
                        # Add comprehensive_report to session state for the Chat mode to use
                        st.session_state.comprehensive_report = report
                    
                    # Show the report
                    st.markdown("## 📋 AI-Powered Analysis Report")
                    
                    # Show the text report
                    # Split the report into sections to insert the pie chart after Model Performance section
                    report_sections = report.split("## 2. Key Business Drivers")
                    
                    if len(report_sections) > 1:
                        # Display first section (Model Performance)
                        st.markdown(report_sections[0])
                        
                        # Display the pie chart after Model Performance section
                        if predictor_data:
                            st.markdown("### 📈 Response Contribution by Media Channel")
                            st.pyplot(pie_chart)
                            
                            # Base info text
                            info_text = """
                            <div class="info-box">
                                <p>This visualization shows the percentage contribution of each media channel to the overall response, 
                                based on the Response Decomposition Waterfall by Predictor from the one-pager model.</p>
                            """                            # Add explanation about filtered negative values if needed
                            if has_negative_values:
                                # Use the helper function to get the note with specific predictor names
                                filtered_note = get_filtered_values_note(has_negative_values)
                                info_text += filtered_note
                            
                            info_text += "</div>"
                            st.markdown(info_text, unsafe_allow_html=True)
                        
                        # Display the rest of the report
                        st.markdown("## 2. Key Business Drivers" + report_sections[1])
                    else:
                        # Fallback if report doesn't have expected structure
                        st.markdown(report)
                        
                        # Display the pie chart if data is available
                        if predictor_data:
                            st.markdown("### 📈 Response Contribution by Media Channel")
                            st.pyplot(pie_chart)
                            
                            # Base info text
                            info_text = """
                            <div class="info-box">
                                <p>This visualization shows the percentage contribution of each media channel to the overall response, 
                                based on the Response Decomposition Waterfall by Predictor from the one-pager model.</p>
                            """                            # Add explanation about filtered negative values if needed
                            if has_negative_values:
                                # Use the helper function to get the note with specific predictor names
                                filtered_note = get_filtered_values_note(has_negative_values)
                                info_text += filtered_note
                            
                            info_text += "</div>"
                            st.markdown(info_text, unsafe_allow_html=True)
                    
                    # Option to download as DOCX
                    doc = Document()
                    
                    # Split the report to insert the pie chart after Model Performance section
                    report_sections = report.split("## 2. Key Business Drivers")
                    
                    if len(report_sections) > 1:
                        # Add first section (Model Performance)
                        add_formatted_text_to_doc(doc, report_sections[0])
                        
                        # Add the pie chart after Model Performance section
                        if predictor_data:
                            negative_predictors = st.session_state.negative_predictors if 'negative_predictors' in st.session_state else None
                            add_chart_to_doc(doc, pie_chart, has_negative_values, negative_predictors=negative_predictors)
                        
                        # Add the rest of the report
                        add_formatted_text_to_doc(doc, "## 2. Key Business Drivers" + report_sections[1])
                    else:
                        # Fallback if report doesn't have expected structure
                        add_formatted_text_to_doc(doc, str(report))
                        
                        # Add the pie chart at the end if structure not recognized
                        if predictor_data:
                            negative_predictors = st.session_state.negative_predictors if 'negative_predictors' in st.session_state else None
                            add_chart_to_doc(doc, pie_chart, has_negative_values, negative_predictors=negative_predictors)
                    
                    # Save temporarily and offer for download
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                        doc.save(temp_file.name)
                        
                        with open(temp_file.name, 'rb') as file:
                            docx_bytes = file.read()
                            
                        st.download_button(
                            label="📄 Download Report as DOCX",
                            data=docx_bytes,
                            file_name="Marketing_Mix_Model_Analysis.docx",
                            help="Download the analysis as a Word document",
                            key="download_generated_report"
                        )
                
                except Exception as e:
                    st.error(f"Error during analysis: {str(e)}")
                    st.markdown("""
                    <div class="error-box">
                        <p>There was an error during analysis. Please ensure:</p>
                        <ul>
                            <li>Both images are clear and high resolution</li>
                            <li>Your OpenAI API key is valid and has sufficient credits</li>
                            <li>Your internet connection is stable</li>
                        </ul>
                    </div>
                    """, unsafe_allow_html=True)
    
    else:  # Chat with Model mode
        # Display chat interface
        st.markdown("""
        <div class="info-box">
            <h3>Chat with Model Mode</h3>
            <p>Ask specific questions about your marketing mix model analysis:</p>
            <ul>
                <li>Which channel has the highest ROI?</li>
                <li>What is the optimal budget allocation?</li>
                <li>How does TV performance compare to Digital?</li>
                <li>What's the R-squared value of this model?</li>
                <li>How do the budget allocation recommendations align with the channel performance?</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        # Check if AI-Powered Analysis has been performed
        has_comprehensive_report = 'comprehensive_report' in st.session_state and st.session_state.comprehensive_report is not None
        
        # Show info message if comprehensive report is not available
        if not has_comprehensive_report:
            st.info("""
            💡 For general questions, you can start chatting right away. 
            
            For detailed marketing mix model analysis, it's recommended to run the AI-Powered Analysis first.
            """)
            
            # Add a button to switch to AI-Powered Analysis mode
            col1, col2 = st.columns([1, 2])
            with col1:
                if st.button("Run AI-Powered Analysis"):
                    # Set a flag that will be checked at the start of the app
                    st.session_state.switch_to_comprehensive = True
                    st.rerun()
        else:
            # Show success message if comprehensive report is available
            st.success("✅ AI-Powered Analysis has been completed! You can now ask specific questions about your marketing mix model.")
        
        # Always show chat interface
        # Display chat history
        for message in st.session_state.chat_history:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        # Chat input
        user_question = st.chat_input(
            "Ask a question about your marketing mix model analysis...",
            disabled=not api_key
        )
        
        # Process chat input when submitted
        if user_question:
            # Add user message to chat
            st.session_state.chat_history.append({"role": "user", "content": user_question})
            
            # Display user message
            with st.chat_message("user"):
                st.markdown(user_question)
            
            # Process with AI
            with st.chat_message("assistant"):
                with st.spinner("Processing your question..."):
                    try:
                        # Use OpenAI client for text-based Q&A
                        client = OpenAI(api_key=api_key)
                        
                        # Create system message
                        system_message = "You are an AI assistant with expertise in marketing mix models and analytics, but you are also capable of answering general knowledge questions accurately. When a question is specifically about marketing mix models, media channels, ROAS, or budgeting, ensure your response is detailed and references the analysis provided. For general inquiries, offer a well-informed answer based on your broad expertise."
                        
                        # Check if comprehensive report is available and include it in the prompt if it is
                        if has_comprehensive_report:
                            comprehensive_report = st.session_state.comprehensive_report
                            prompt = f"""
                            I have the following comprehensive analysis of marketing mix model images:
                            
                            {comprehensive_report}
                            
                            User's question:
                            {user_question}
                            
                            Instructions:
                            - If the question relates specifically to marketing mix models (including topics such as media channels, ROAS, budgeting, etc.) or any content from the analysis above, generate a detailed, data-driven response that directly references the provided analysis.
                            - If the question is of a general nature (e.g., "What is machine learning?" or any topic not covered by the analysis), provide a clear, accurate answer using your general expertise.

                            Ensure that marketing-specific responses make use of the analysis data, while general knowledge responses rely on accurate, broad-based insights.
                            """
                        else:
                            # No comprehensive report available - just answer the question as best as possible
                            prompt = f"""
                            User's question:
                            {user_question}
                            
                            Instructions:
                            - If the question relates specifically to marketing mix models (including topics such as media channels, ROAS, budgeting, etc.), explain that no marketing mix model analysis has been run yet, but provide general information about the topic.
                            - If the question is of a general nature (e.g., "What is machine learning?" or any topic not specifically about marketing mix models), provide a clear, accurate answer using your general expertise.
                            
                            Note that the user hasn't run an AI-Powered Analysis of their marketing mix model yet, so you don't have specific data about their model to reference.
                            """
                        
                        # Send to the API
                        response = client.chat.completions.create(
                            model="gpt-4o",
                            temperature=0.0,
                            messages=[
                                {"role": "system", "content": system_message},
                                {"role": "user", "content": prompt}
                            ],
                            max_tokens=1500
                        )
                        
                        # Get and display the response
                        answer = response.choices[0].message.content
                        st.markdown(answer)
                        
                        # Add assistant message to chat
                        st.session_state.chat_history.append({"role": "assistant", "content": answer})
                    
                    except Exception as e:
                        error_msg = f"Error processing your question: {str(e)}"
                        st.error(error_msg)
                        st.session_state.chat_history.append({"role": "assistant", "content": error_msg})
        
        # Clear chat button
        if st.button("Clear Chat History", help="Click to clear the current conversation"):
            st.session_state.chat_history = []
            st.rerun()

# Run the app
if __name__ == "__main__":
    main()
